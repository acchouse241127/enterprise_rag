"""
Conversation service for storing, exporting, and sharing conversations.

Phase 3.3: 对话导出与分享
Author: C2
Date: 2026-02-14
"""

import io
import uuid
from datetime import datetime, timedelta
from typing import Literal

from sqlalchemy import select, desc
from sqlalchemy.orm import Session

from app.models.conversation import Conversation, ConversationMessage


class ConversationService:
    """Service for managing conversations."""

    @staticmethod
    def create_conversation(
        db: Session,
        knowledge_base_id: int,
        user_id: int | None = None,
        title: str | None = None,
        conversation_id: str | None = None,
    ) -> Conversation:
        """Create a new conversation. conversation_id 可选，不传则生成 UUID（OPT-024 支持前端传入）。"""
        conversation = Conversation(
            conversation_id=(conversation_id and conversation_id.strip()) or str(uuid.uuid4()),
            knowledge_base_id=knowledge_base_id,
            user_id=user_id,
            title=title,
        )
        db.add(conversation)
        db.commit()
        db.refresh(conversation)
        return conversation

    @staticmethod
    def get_conversation(db: Session, conversation_id: int) -> Conversation | None:
        """Get conversation by ID."""
        stmt = select(Conversation).where(Conversation.id == conversation_id)
        return db.execute(stmt).scalar_one_or_none()

    @staticmethod
    def get_by_conversation_id(db: Session, conversation_id: str) -> Conversation | None:
        """Get conversation by UUID string."""
        stmt = select(Conversation).where(Conversation.conversation_id == conversation_id)
        return db.execute(stmt).scalar_one_or_none()

    @staticmethod
    def get_by_share_token(db: Session, share_token: str) -> Conversation | None:
        """Get conversation by share token."""
        stmt = select(Conversation).where(Conversation.share_token == share_token)
        conv = db.execute(stmt).scalar_one_or_none()
        if conv and conv.is_shared:
            # Check expiration
            if conv.share_expires_at and conv.share_expires_at < datetime.now():
                return None
            return conv
        return None

    @staticmethod
    def list_conversations(
        db: Session,
        user_id: int | None = None,
        knowledge_base_id: int | None = None,
        limit: int = 50,
        offset: int = 0,
    ) -> list[Conversation]:
        """List conversations with optional filters."""
        stmt = select(Conversation).order_by(desc(Conversation.updated_at))

        if user_id:
            stmt = stmt.where(Conversation.user_id == user_id)
        if knowledge_base_id:
            stmt = stmt.where(Conversation.knowledge_base_id == knowledge_base_id)

        stmt = stmt.limit(limit).offset(offset)
        return list(db.execute(stmt).scalars().all())

    @staticmethod
    def persist_qa_turn(
        db: Session,
        conversation_id_str: str,
        knowledge_base_id: int,
        user_id: int | None,
        question: str,
        answer: str,
    ) -> None:
        """OPT-024: 获取或创建会话并写入本轮 user/assistant 消息，供 RAG 问答与对话管理打通。"""
        if not (conversation_id_str and conversation_id_str.strip()):
            return
        conv = ConversationService.get_by_conversation_id(db, conversation_id_str.strip())
        if conv is None:
            conv = ConversationService.create_conversation(
                db,
                knowledge_base_id=knowledge_base_id,
                user_id=user_id,
                title=(question[:50] + "..." if len(question) > 50 else question) or "新对话",
                conversation_id=conversation_id_str.strip(),
            )
        ConversationService.add_message(db, conv.id, "user", question)
        ConversationService.add_message(db, conv.id, "assistant", answer)
        return

    @staticmethod
    def persist_qa_turn_standalone(
        conversation_id_str: str,
        knowledge_base_id: int,
        user_id: int | None,
        question: str,
        answer: str,
    ) -> None:
        """OPT-024: 独立 session 持久化一轮 QA（用于 stream 结束后，无法使用请求 scope 的 db 时）。"""
        from app.core.database import SessionLocal
        db = SessionLocal()
        try:
            ConversationService.persist_qa_turn(
                db, conversation_id_str, knowledge_base_id, user_id, question, answer
            )
        finally:
            db.close()

    @staticmethod
    def add_message(
        db: Session,
        conversation_id: int,
        role: Literal["user", "assistant"],
        content: str,
        extra_data: dict | None = None,
    ) -> ConversationMessage:
        """Add a message to a conversation."""
        message = ConversationMessage(
            conversation_id=conversation_id,
            role=role,
            content=content,
            extra_data=extra_data,
        )
        db.add(message)
        
        # Update conversation timestamp
        conv = ConversationService.get_conversation(db, conversation_id)
        if conv:
            conv.updated_at = datetime.now()
        
        db.commit()
        db.refresh(message)
        return message

    @staticmethod
    def get_messages(db: Session, conversation_id: int) -> list[ConversationMessage]:
        """Get all messages in a conversation."""
        stmt = (
            select(ConversationMessage)
            .where(ConversationMessage.conversation_id == conversation_id)
            .order_by(ConversationMessage.created_at)
        )
        return list(db.execute(stmt).scalars().all())

    @staticmethod
    def update_title(db: Session, conversation_id: int, title: str) -> Conversation | None:
        """Update conversation title."""
        conv = ConversationService.get_conversation(db, conversation_id)
        if not conv:
            return None
        conv.title = title
        db.commit()
        db.refresh(conv)
        return conv

    @staticmethod
    def enable_sharing(
        db: Session,
        conversation_id: int,
        expires_in_days: int | None = 7,
    ) -> tuple[Conversation | None, str | None]:
        """Enable sharing for a conversation."""
        conv = ConversationService.get_conversation(db, conversation_id)
        if not conv:
            return None, None

        if not conv.share_token:
            conv.share_token = Conversation.generate_share_token()
        conv.is_shared = True
        if expires_in_days:
            conv.share_expires_at = datetime.now() + timedelta(days=expires_in_days)
        else:
            conv.share_expires_at = None

        db.commit()
        db.refresh(conv)
        return conv, conv.share_token

    @staticmethod
    def disable_sharing(db: Session, conversation_id: int) -> Conversation | None:
        """Disable sharing for a conversation."""
        conv = ConversationService.get_conversation(db, conversation_id)
        if not conv:
            return None
        conv.is_shared = False
        db.commit()
        db.refresh(conv)
        return conv

    @staticmethod
    def export_to_markdown(db: Session, conversation_id: int) -> str | None:
        """Export conversation to Markdown format."""
        conv = ConversationService.get_conversation(db, conversation_id)
        if not conv:
            return None

        messages = ConversationService.get_messages(db, conversation_id)
        
        lines = []
        lines.append(f"# {conv.title or '对话记录'}")
        lines.append("")
        lines.append(f"**导出时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        lines.append(f"**对话ID**: {conv.conversation_id}")
        lines.append("")
        lines.append("---")
        lines.append("")

        for msg in messages:
            role_label = "**用户**" if msg.role == "user" else "**助手**"
            lines.append(f"### {role_label}")
            lines.append("")
            lines.append(msg.content)
            lines.append("")
            
            # Add citations if present
            if msg.extra_data and msg.extra_data.get("citations"):
                lines.append("**引用来源**:")
                for citation in msg.extra_data["citations"]:
                    lines.append(f"- {citation.get('document_title', '未知文档')}")
                lines.append("")
            
            lines.append("---")
            lines.append("")

        return "\n".join(lines)

    @staticmethod
    def _register_cjk_font() -> str:
        """Register a CJK font from system; return font name for use, or 'Helvetica'."""
        import os
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        candidates = [
            "C:/Windows/Fonts/simhei.ttf",
            "C:/Windows/Fonts/msyh.ttc",
            "C:/Windows/Fonts/simsun.ttc",
            "C:/Windows/Fonts/msyhbd.ttc",
            "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
            "/System/Library/Fonts/PingFang.ttc",
        ]
        for path in candidates:
            if not os.path.isfile(path):
                continue
            try:
                pdfmetrics.registerFont(TTFont("CJK", path))
                return "CJK"
            except Exception:
                continue
        return "Helvetica"

    @staticmethod
    def export_to_pdf_bytes(db: Session, conversation_id: int) -> bytes | None:
        """
        Export conversation to PDF. Uses a CJK font from system path when available
        to avoid black squares for Chinese text.
        """
        conv = ConversationService.get_conversation(db, conversation_id)
        if not conv:
            return None

        messages = ConversationService.get_messages(db, conversation_id)

        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.units import cm
        except ImportError:
            md_content = ConversationService.export_to_markdown(db, conversation_id)
            return md_content.encode("utf-8") if md_content else None

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm)
        styles = getSampleStyleSheet()
        font_name = ConversationService._register_cjk_font()

        chinese_style = ParagraphStyle(
            "Chinese",
            parent=styles["Normal"],
            fontName=font_name,
            fontSize=10,
            leading=14,
        )
        title_style = ParagraphStyle(
            "Title",
            parent=styles["Heading1"],
            fontName=font_name,
            fontSize=16,
            spaceAfter=20,
        )

        story = []
        story.append(Paragraph((conv.title or "对话记录").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), title_style))
        story.append(Spacer(1, 0.5*cm))
        meta_text = f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        story.append(Paragraph(meta_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), chinese_style))
        story.append(Spacer(1, 1*cm))

        for msg in messages:
            role_label = "用户:" if msg.role == "user" else "助手:"
            story.append(Paragraph(f"<b>{role_label}</b>", chinese_style))
            story.append(Spacer(1, 0.2*cm))
            content = (msg.content or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")
            story.append(Paragraph(content, chinese_style))
            story.append(Spacer(1, 0.5*cm))

        doc.build(story)
        return buffer.getvalue()

    @staticmethod
    def export_to_docx_bytes(db: Session, conversation_id: int) -> bytes | None:
        """Export conversation to DOCX. Returns bytes or None."""
        conv = ConversationService.get_conversation(db, conversation_id)
        if not conv:
            return None
        messages = ConversationService.get_messages(db, conversation_id)
        try:
            from docx import Document as DocxDocument
        except ImportError:
            return None
        doc = DocxDocument()
        doc.add_heading(conv.title or "对话记录", 0)
        doc.add_paragraph(f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"对话ID: {conv.conversation_id}")
        doc.add_paragraph("")
        for msg in messages:
            role_label = "用户" if msg.role == "user" else "助手"
            p = doc.add_paragraph()
            run = p.add_run(f"{role_label}: ")
            run.bold = True
            p.add_run(msg.content or "")
            if msg.extra_data and msg.extra_data.get("citations"):
                p = doc.add_paragraph()
                p.add_run("引用来源: ").bold = True
                for c in msg.extra_data["citations"]:
                    doc.add_paragraph(f"  - {c.get('document_title', '未知文档')}", style="List Bullet")
            doc.add_paragraph("")
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    @staticmethod
    def delete_conversation(db: Session, conversation_id: int) -> bool:
        """Delete a conversation and all its messages."""
        conv = ConversationService.get_conversation(db, conversation_id)
        if not conv:
            return False
        db.delete(conv)
        db.commit()
        return True
